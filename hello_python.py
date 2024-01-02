'''
Created on Dec 23, 2023

@author: RJM
'''
print('Hello World')

from docx import Document

document = Document()

heading1 = document.add_heading('Cunning Strike')

heading2 = document.add_heading('Instant')

paragraph1 = document.add_paragraph('Cunning Strike deals 2 damage to target creature and 2 damage to target player or planeswalker.')

paragraph2 = document.add_paragraph('Draw a card.')

paragraph3 = document.add_paragraph('“The opponent who blocks the path, becomes the path.” —Shu Yun, the Silent Tempest')

document.save('demo.docx')